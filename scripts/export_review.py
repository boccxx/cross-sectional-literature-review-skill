#!/usr/bin/env python3
"""Export literature_review_synthesis.md to LaTeX (.tex), Word (.docx), and PDF.

Export strategy (in order of preference):
  1. LaTeX (.tex): pandoc if available
  2. Word (.docx): pandoc if available, otherwise python-docx with built-in markdown parsing
  3. PDF: pandoc with --pdf-engine (xelatex, pdflatex, or weasyprint), otherwise
     instructs the user to print the DOCX to PDF

Requirements:
  - LaTeX via pandoc: `brew install pandoc`
  - Word via pandoc: `brew install pandoc` or https://pandoc.org/installing.html
  - Word via python-docx: `pip install python-docx`
  - PDF via pandoc+xelatex: `brew install pandoc` + a LaTeX distribution (MacTeX/MiKTeX)
  - PDF via pandoc+weasyprint: `pip install weasyprint` (no LaTeX needed)

Usage:
    python export_review.py \\
        --input literature_review_synthesis.md \\
        --output-dir ./literature_review \\
        [--title "Your Review Title"] \\
        [--author "Author Name"] \\
        [--lang zh-CN] [--mainfont "Songti SC"] \\
        [--no-pdf] [--no-word] [--no-tex] \\
        [--pdf-engine xelatex]

Outputs:
    literature_review.tex
    literature_review.docx
    literature_review.pdf (if possible)
"""
from __future__ import annotations

import argparse
import hashlib
import re
import subprocess
import sys
import textwrap
from pathlib import Path


# ---------------------------------------------------------------------------
# Pandoc helpers
# ---------------------------------------------------------------------------

def find_pandoc() -> str | None:
    """Return the path to pandoc if available, else None."""
    for cmd in ("pandoc", "/usr/local/bin/pandoc", "/opt/homebrew/bin/pandoc"):
        try:
            result = subprocess.run([cmd, "--version"], capture_output=True, timeout=10)
            if result.returncode == 0:
                return cmd
        except (FileNotFoundError, subprocess.TimeoutExpired):
            continue
    return None


def default_mainfont(lang: str) -> str:
    lowered = (lang or "").lower()
    if lowered.startswith("zh"):
        return "Songti SC"
    if lowered.startswith("ja"):
        return "Hiragino Mincho ProN"
    if lowered.startswith("ko"):
        return "AppleMyungjo"
    return "Times New Roman"


def bind_docx_source_hash(path: Path, source_sha256: str) -> bool:
    """Embed the source digest in standard DOCX core properties."""
    try:
        from docx import Document  # type: ignore

        doc = Document(str(path))
        doc.core_properties.keywords = f"source_sha256:{source_sha256}"
        doc.save(str(path))
        return True
    except Exception as exc:
        print(f"  [docx source binding error] {exc}")
        return False


def pandoc_to_docx(pandoc: str, md_path: Path, out_path: Path, title: str, author: str, lang: str, source_sha256: str) -> bool:
    cmd = [
        pandoc,
        str(md_path),
        "--from", "markdown+smart",
        "--to", "docx",
        "--output", str(out_path),
        "--standalone",
        f"--metadata=title:{title}",
    ]
    if author:
        cmd.append(f"--metadata=author:{author}")
    if lang:
        cmd.append(f"--metadata=lang:{lang}")
    try:
        result = subprocess.run(cmd, capture_output=True, text=True, timeout=120)
        if result.returncode == 0:
            return bind_docx_source_hash(out_path, source_sha256)
        print(f"  [pandoc docx error] {result.stderr[:400]}")
        return False
    except Exception as exc:
        print(f"  [pandoc docx exception] {exc}")
        return False


def pandoc_to_pdf(
    pandoc: str,
    md_path: Path,
    out_path: Path,
    title: str,
    author: str,
    pdf_engine: str,
    lang: str,
    mainfont: str,
    source_sha256: str,
) -> bool:
    cmd = [
        pandoc,
        str(md_path),
        "--from", "markdown+smart",
        "--to", "pdf",
        "--output", str(out_path),
        "--standalone",
        f"--pdf-engine={pdf_engine}",
        f"--metadata=title:{title}",
        f"--metadata=keywords:source_sha256_{source_sha256}",
        "--variable", "geometry:margin=2.5cm",
        "--variable", "fontsize=11pt",
        f"--variable=mainfont:{mainfont}",
    ]
    if author:
        cmd.append(f"--metadata=author:{author}")
    if lang:
        cmd.append(f"--metadata=lang:{lang}")
    try:
        result = subprocess.run(cmd, capture_output=True, text=True, timeout=300)
        if result.returncode == 0:
            return True
        print(f"  [pandoc pdf error] {result.stderr[:400]}")
        return False
    except Exception as exc:
        print(f"  [pandoc pdf exception] {exc}")
        return False


def pandoc_to_html(pandoc: str, md_path: Path, out_path: Path, title: str, author: str, lang: str) -> bool:
    """Intermediate HTML for weasyprint-based PDF conversion."""
    cmd = [
        pandoc,
        str(md_path),
        "--from", "markdown+smart",
        "--to", "html5",
        "--output", str(out_path),
        "--standalone",
        "--embed-resources",
        f"--metadata=title:{title}",
    ]
    if author:
        cmd.append(f"--metadata=author:{author}")
    if lang:
        cmd.append(f"--metadata=lang:{lang}")
    try:
        result = subprocess.run(cmd, capture_output=True, text=True, timeout=120)
        return result.returncode == 0
    except Exception:
        return False


def weasyprint_from_html(html_path: Path, pdf_path: Path) -> bool:
    try:
        import weasyprint  # type: ignore
        weasyprint.HTML(filename=str(html_path)).write_pdf(str(pdf_path))
        return True
    except ImportError:
        print("  [weasyprint] Not installed. Run: pip install weasyprint")
        return False
    except Exception as exc:
        print(f"  [weasyprint error] {exc}")
        return False


# ---------------------------------------------------------------------------
# python-docx fallback
# ---------------------------------------------------------------------------

def _apply_inline(para, text: str) -> None:
    """Apply bold/italic inline formatting to a paragraph."""
    # Split on **bold** and *italic* patterns
    pattern = re.compile(r"(\*\*(.+?)\*\*|\*(.+?)\*|`(.+?)`)")
    pos = 0
    for m in pattern.finditer(text):
        if m.start() > pos:
            para.add_run(text[pos: m.start()])
        full = m.group(0)
        if full.startswith("**"):
            run = para.add_run(m.group(2))
            run.bold = True
        elif full.startswith("`"):
            run = para.add_run(m.group(4))
            run.font.name = "Courier New"
        else:
            run = para.add_run(m.group(3))
            run.italic = True
        pos = m.end()
    if pos < len(text):
        para.add_run(text[pos:])


def python_docx_export(md_path: Path, out_path: Path, title: str, author: str, source_sha256: str) -> bool:
    try:
        from docx import Document  # type: ignore
        from docx.shared import Pt, RGBColor  # type: ignore
        from docx.enum.text import WD_ALIGN_PARAGRAPH  # type: ignore
    except ImportError:
        print("  [python-docx] Not installed. Run: pip install python-docx")
        return False

    doc = Document()
    doc.core_properties.keywords = f"source_sha256:{source_sha256}"

    # Styles
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(11)

    # Cover info
    if title:
        t = doc.add_heading(title, level=0)
        t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if author:
        a = doc.add_paragraph(author)
        a.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if title or author:
        doc.add_paragraph("")

    md_text = md_path.read_text(encoding="utf-8")
    lines = md_text.splitlines()

    i = 0
    while i < len(lines):
        line = lines[i]

        # Heading
        heading_match = re.match(r"^(#{1,4})\s+(.*)", line)
        if heading_match:
            level = len(heading_match.group(1))
            heading_text = heading_match.group(2).strip()
            # Strip markdown formatting from heading
            heading_clean = re.sub(r"\*\*(.+?)\*\*", r"\1", heading_text)
            heading_clean = re.sub(r"\*(.+?)\*", r"\1", heading_clean)
            doc.add_heading(heading_clean, level=min(level, 4))
            i += 1
            continue

        # Horizontal rule
        if re.match(r"^---+$", line.strip()):
            doc.add_paragraph("─" * 40)
            i += 1
            continue

        # Blank line
        if not line.strip():
            i += 1
            continue

        # Table (simplified: detect | rows)
        if line.strip().startswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i])
                i += 1
            # Parse table
            rows = []
            for tl in table_lines:
                if re.match(r"^\|[-:| ]+\|$", tl.strip()):
                    continue
                cells = [c.strip() for c in tl.strip().strip("|").split("|")]
                rows.append(cells)
            if rows:
                n_cols = max(len(r) for r in rows)
                table = doc.add_table(rows=len(rows), cols=n_cols)
                table.style = "Table Grid"
                for ri, row in enumerate(rows):
                    for ci, cell_text in enumerate(row):
                        if ci < n_cols:
                            cell_clean = re.sub(r"\*\*(.+?)\*\*", r"\1", cell_text)
                            cell_clean = re.sub(r"\*(.+?)\*", r"\1", cell_clean)
                            cell_clean = re.sub(r"`(.+?)`", r"\1", cell_clean)
                            table.cell(ri, ci).text = cell_clean
            continue

        # Bullet list
        if re.match(r"^[-*+]\s+", line):
            bullet_text = re.sub(r"^[-*+]\s+", "", line)
            p = doc.add_paragraph(style="List Bullet")
            _apply_inline(p, bullet_text)
            i += 1
            continue

        # Numbered list
        if re.match(r"^\d+\.\s+", line):
            num_text = re.sub(r"^\d+\.\s+", "", line)
            p = doc.add_paragraph(style="List Number")
            _apply_inline(p, num_text)
            i += 1
            continue

        # Blockquote
        if line.startswith(">"):
            quote_text = re.sub(r"^>\s*", "", line)
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Pt(24)
            run = p.add_run(quote_text)
            run.italic = True
            i += 1
            continue

        # Normal paragraph (collect continuation lines)
        para_lines = [line]
        i += 1
        while i < len(lines) and lines[i].strip() and not re.match(r"^#{1,4}\s", lines[i]) and not lines[i].startswith("|"):
            para_lines.append(lines[i])
            i += 1
        para_text = " ".join(para_lines)
        p = doc.add_paragraph()
        _apply_inline(p, para_text)

    doc.save(str(out_path))
    return True


def pandoc_to_tex(pandoc: str, md_path: Path, out_path: Path, title: str, author: str, lang: str, source_sha256: str) -> bool:
    cmd = [
        pandoc,
        str(md_path),
        "--from", "markdown+smart",
        "--to", "latex",
        "--output", str(out_path),
        "--standalone",
        f"--metadata=title:{title}",
    ]
    if author:
        cmd.append(f"--metadata=author:{author}")
    if lang:
        cmd.append(f"--metadata=lang:{lang}")
    try:
        result = subprocess.run(cmd, capture_output=True, text=True, timeout=120)
        if result.returncode == 0:
            text = out_path.read_text(encoding="utf-8")
            out_path.write_text(f"% source_sha256:{source_sha256}\n{text}", encoding="utf-8")
            return True
        print(f"  [pandoc tex error] {result.stderr[:400]}")
        return False
    except Exception as exc:
        print(f"  [pandoc tex exception] {exc}")
        return False


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def infer_title_from_md(md_path: Path) -> str:
    """Extract the first H1 heading from the markdown file."""
    try:
        for line in md_path.read_text(encoding="utf-8").splitlines():
            m = re.match(r"^#\s+(.+)", line)
            if m:
                return m.group(1).strip()
    except Exception:
        pass
    return md_path.stem.replace("_", " ").title()


def main() -> None:
    parser = argparse.ArgumentParser(
        description="Export literature_review_synthesis.md to LaTeX (.tex), Word (.docx), and PDF.",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog=textwrap.dedent("""\
            Requirements:
              LaTeX via pandoc: brew install pandoc
              Word via pandoc: brew install pandoc
              Word via python-docx: pip install python-docx
              PDF via pandoc+xelatex: pandoc + MacTeX / MiKTeX
              PDF via pandoc+weasyprint: pandoc + pip install weasyprint
        """),
    )
    parser.add_argument("--input", required=True, help="Path to literature_review_synthesis.md (or any Markdown file)")
    parser.add_argument("--output-dir", required=True, help="Directory to write .docx and .pdf outputs")
    parser.add_argument("--title", default="", help="Document title (auto-detected from first H1 if omitted)")
    parser.add_argument("--author", default="", help="Author name for document metadata")
    parser.add_argument("--stem", default="literature_review", help="Output file stem (default: literature_review)")
    parser.add_argument("--lang", default="en", help="Document language tag, e.g. en, zh-CN, ja, de, fr, es")
    parser.add_argument("--mainfont", default="", help="Main font for PDF export; auto-selected by language if omitted")
    parser.add_argument("--no-tex", action="store_true", help="Skip LaTeX export")
    parser.add_argument("--no-word", action="store_true", help="Skip Word export")
    parser.add_argument("--no-pdf", action="store_true", help="Skip PDF export")
    parser.add_argument(
        "--pdf-engine",
        default="",
        help="Pandoc PDF engine: xelatex (default), pdflatex, lualatex, weasyprint",
    )
    parser.add_argument("--insecure-skip-verify", action="store_true", help="(Unused; reserved for consistency)")
    args = parser.parse_args()

    md_path = Path(args.input)
    if not md_path.exists():
        print(f"[export_review] Input file not found: {md_path}")
        sys.exit(1)

    output_dir = Path(args.output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)

    title = args.title or infer_title_from_md(md_path)
    author = args.author
    stem = args.stem
    lang = args.lang
    mainfont = args.mainfont or default_mainfont(lang)
    source_sha256 = hashlib.sha256(md_path.read_bytes()).hexdigest()

    tex_path = output_dir / f"{stem}.tex"
    docx_path = output_dir / f"{stem}.docx"
    pdf_path = output_dir / f"{stem}.pdf"

    pandoc = find_pandoc()
    if pandoc:
        print(f"[export_review] pandoc found: {pandoc}")
    else:
        print("[export_review] pandoc not found; will use python-docx fallback for Word")

    # ---- LaTeX export ----
    if not args.no_tex:
        print(f"[export_review] Exporting LaTeX → {tex_path}")
        if pandoc and pandoc_to_tex(pandoc, md_path, tex_path, title, author, lang, source_sha256):
            print(f"[export_review] LaTeX saved via pandoc: {tex_path}")
        else:
            print("[export_review] LaTeX export failed. Install pandoc to enable .tex export.")

    # ---- Word export ----
    if not args.no_word:
        print(f"[export_review] Exporting Word → {docx_path}")
        if pandoc and pandoc_to_docx(pandoc, md_path, docx_path, title, author, lang, source_sha256):
            print(f"[export_review] Word saved via pandoc: {docx_path}")
        elif python_docx_export(md_path, docx_path, title, author, source_sha256):
            print(f"[export_review] Word saved via python-docx: {docx_path}")
        else:
            print("[export_review] Word export failed. Install pandoc or python-docx.")

    # ---- PDF export ----
    if not args.no_pdf:
        print(f"[export_review] Exporting PDF → {pdf_path}")
        pdf_ok = False

        if pandoc:
            engines_to_try = []
            if args.pdf_engine:
                engines_to_try = [args.pdf_engine]
            else:
                engines_to_try = ["xelatex", "pdflatex", "lualatex", "weasyprint"]

            for engine in engines_to_try:
                print(f"  [export_review] Trying pdf-engine: {engine}")
                if engine == "weasyprint":
                    html_tmp = output_dir / f"{stem}_tmp.html"
                    if pandoc_to_html(pandoc, md_path, html_tmp, title, author, lang):
                        pdf_ok = weasyprint_from_html(html_tmp, pdf_path)
                        try:
                            html_tmp.unlink()
                        except Exception:
                            pass
                    if pdf_ok:
                        break
                else:
                    pdf_ok = pandoc_to_pdf(pandoc, md_path, pdf_path, title, author, engine, lang, mainfont, source_sha256)
                    if pdf_ok:
                        break

        if pdf_ok:
            print(f"[export_review] PDF saved: {pdf_path}")
        else:
            print("[export_review] PDF export failed. Options:")
            print("  1. Install pandoc + a LaTeX distribution (MacTeX/MiKTeX)")
            print("  2. Install pandoc + weasyprint: pip install weasyprint")
            print("  3. Open the .docx in Word or LibreOffice and export to PDF")

    print("\n[export_review] Done.")


if __name__ == "__main__":
    main()
